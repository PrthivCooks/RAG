import sys
import pickle
import numpy as np
import faiss
import spacy
from pathlib import Path
from tqdm import tqdm
from rank_bm25 import BM25Okapi

from config import (
    DATA_DIR,
    INDEX_DIR,
    EMBEDDING_MODEL_NAME,
    SPACY_MODEL_NAME,
    DEFAULT_CHUNK_SIZE,
    DEFAULT_CHUNK_OVERLAP
)
from core.document_processor import DocumentProcessor
from core.graph_builder import GraphBuilder

def check_and_load_spacy():
    """Load spaCy model, downloading it if not present."""
    try:
        print(f"Loading spaCy model: {SPACY_MODEL_NAME}...")
        return spacy.load(SPACY_MODEL_NAME)
    except OSError:
        print(f"spaCy model '{SPACY_MODEL_NAME}' not found. Downloading it...")
        try:
            from spacy.cli import download
            download(SPACY_MODEL_NAME)
            return spacy.load(SPACY_MODEL_NAME)
        except Exception as e:
            print(f"Failed to download spaCy model: {e}")
            print("Please run 'python -m spacy download en_core_web_sm' manually with elevated permissions if needed.")
            sys.exit(1)

def run_ingestion():
    """
    Main ingestion execution pipeline.
    Loads documents, chunks them semantically, builds FAISS, BM25, and Graph indexes,
    and serializes them to disk.
    """
    # Ensure folders exist
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    INDEX_DIR.mkdir(parents=True, exist_ok=True)

    # 1. Load spaCy
    nlp = check_and_load_spacy()

    # 2. Load Embedding Model
    print(f"Loading SentenceTransformer embedding model: {EMBEDDING_MODEL_NAME}...")
    try:
        from sentence_transformers import SentenceTransformer
        embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    except Exception as e:
        print(f"Error loading embedding model: {e}")
        sys.exit(1)

    # Check if data directory has files
    files = list(DATA_DIR.glob("*"))
    supported_files = [f for f in files if f.suffix.lower() in [".pdf", ".docx", ".txt"]]

    if not supported_files:
        print("\n" + "="*80)
        print(f"WARNING: No documents found in '{DATA_DIR.absolute()}'")
        print("Please place some PDF, DOCX, or TXT files in that folder or upload them via the Streamlit UI.")
        print("="*80 + "\n")
        return

    # 3. Process documents to chunks
    print("Processing documents...")
    processor = DocumentProcessor(embedding_model, nlp)
    
    # Run ingestion
    chunks = processor.process_directory(
        directory_path=DATA_DIR,
        chunk_size=DEFAULT_CHUNK_SIZE,
        chunk_overlap=DEFAULT_CHUNK_OVERLAP,
        use_semantic=True
    )
    
    if not chunks:
        print("No document text was successfully extracted and chunked. Ingestion aborted.")
        return

    print(f"Successfully generated {len(chunks)} chunks from {len(set(c['doc_id'] for c in chunks))} documents.")

    # 4. Save Chunk Details Lookup
    chunks_lookup = {c["chunk_id"]: c for c in chunks}
    chunks_path = INDEX_DIR / "chunks.pkl"
    with open(chunks_path, "wb") as f:
        pickle.dump(chunks_lookup, f)
    print(f"Saved chunk lookup file (size: {len(chunks_lookup)} entries) to {chunks_path}")

    # 5. Build and Save Sparse Index (BM25)
    print("Building BM25 Sparse Index...")
    tokenized_corpus = []
    chunk_ids = []
    
    for c in tqdm(chunks, desc="Tokenizing Chunks for BM25"):
        # Simple lowercase tokenization of text
        tokens = [t.text.lower() for t in nlp(c["text"])]
        tokenized_corpus.append(tokens)
        chunk_ids.append(c["chunk_id"])
        
    bm25 = BM25Okapi(tokenized_corpus)
    bm25_path = INDEX_DIR / "bm25.pkl"
    with open(bm25_path, "wb") as f:
        pickle.dump((bm25, chunk_ids), f)
    print(f"Saved BM25 index to {bm25_path}")

    # 6. Build and Save Dense Index (FAISS)
    print("Building FAISS Dense Index...")
    chunk_texts = [c["text"] for c in chunks]
    
    # We compute chunk embeddings
    print("Generating dense embeddings for chunks...")
    embeddings = embedding_model.encode(chunk_texts, show_progress_bar=True, convert_to_numpy=True)
    
    # Normalize embeddings for cosine similarity
    faiss.normalize_L2(embeddings)
    dimension = embeddings.shape[1]
    
    # Create flat inner product index (equivalent to cosine similarity on normalized vectors)
    faiss_index = faiss.IndexFlatIP(dimension)
    faiss_index.add(embeddings)
    
    faiss_path = INDEX_DIR / "faiss.idx"
    faiss.write_index(faiss_index, str(faiss_path))
    print(f"Saved FAISS index ({embeddings.shape[0]} vectors) to {faiss_path}")

    # 7. Build and Save Knowledge Graph
    print("Building Knowledge Graph Index...")
    graph_builder = GraphBuilder(nlp)
    
    # Feed chunks to GraphBuilder
    for c in tqdm(chunks, desc="Extracting Entities and Relations for Graph"):
        graph_builder.add_chunk_to_graph(c["chunk_id"], c["text"], c["metadata"])
        
    graph_path = INDEX_DIR / "graph.pkl"
    graph_builder.save_index(graph_path)
    print("Ingestion completed successfully.")

def create_sample_documents():
    """Generates standard corporate policy files for search demonstration."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    
    # Document 1: Data Retention
    retention_content = """
    Department: Security
    Date: 2026-03-15
    Title: Enterprise Data Retention Policy
    
    1. PURPOSE
    This policy defines the retention periods for all corporate data stored on enterprise servers. It is designed to ensure compliance with legal obligations and to optimize storage costs.
    
    2. RETENTION PERIODS
    - Financial Records: All tax filings, audit trails, billing invoices, and general ledger reports must be retained for exactly 7 years.
    - Customer Account Records: Retained for 3 years following the closure of the account.
    - Employee Records: Retained for 5 years after termination of employment.
    - Marketing Materials: Retained for 1 year after the campaign concludes.
    - Email Correspondence: System logs and emails are archived for 2 years and then permanently deleted unless flagged for legal hold.
    
    3. DATA DELETION
    Once the retention period has expired, data must be securely erased using certified shredding protocols. For digital files, files must be overwritten with random data before deletion.
    
    4. GDPR COMPLIANCE
    In accordance with GDPR compliance requirements, personal data must not be stored longer than necessary for the purpose it was gathered. European citizens can request data erasure (Right to be Forgotten) at any time.
    """
    with open(DATA_DIR / "data_retention_policy.txt", "w", encoding="utf-8") as f:
        f.write(retention_content.strip())
        
    # Document 2: Refund Policy (DOCX format)
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("Customer Refund and Dispute Policy", 0)
        
        # Add metadata block
        p_meta = doc.add_paragraph()
        p_meta.add_run("Department: Finance\n").bold = True
        p_meta.add_run("Date: 2026-04-01\n").bold = True
        p_meta.add_run("Version: 2.4\n").italic = True
        
        doc.add_heading("1. Refund Criteria", level=1)
        doc.add_paragraph(
            "Customer refund requests are eligible for full reimbursement under the following circumstances:\n"
            "a) The billing dispute was filed within 30 days of the invoice date.\n"
            "b) The product or service was not delivered or was defective.\n"
            "c) Double billing occurred due to a processing system error."
        )
        
        doc.add_heading("2. How to handle customer refund requests", level=1)
        doc.add_paragraph(
            "When a support agent receives a refund request, they must follow these steps:\n"
            "1. Verify the customer invoice in the billing system.\n"
            "2. Ensure the customer accounts are active and have no prior fraudulent flags.\n"
            "3. If the request is under $100, the support agent can approve it instantly.\n"
            "4. If the refund is over $100, escalate to the Finance Manager for written approval.\n"
            "5. Process refund through Stripe or the original credit card payment gateway."
        )
        
        doc.add_heading("3. Compliance", level=1)
        doc.add_paragraph(
            "All approved refunds must be logged in the quarterly audit ledger. Standard processing times are 5-10 business days depending on the customer's banking institution."
        )
        doc.save(str(DATA_DIR / "customer_refund_policy.docx"))
    except Exception as e:
        print(f"Could not generate docx sample: {e}. Writing txt fallback.")
        refund_fallback = """
        Department: Finance
        Date: 2026-04-01
        Title: Customer Refund and Dispute Policy
        
        1. Refund Criteria:
        Customer refund requests are eligible for full reimbursement if a billing dispute is filed within 30 days of the invoice date, the product was defective, or double billing occurred.
        
        2. How to handle customer refund requests:
        - Verify customer invoice in the billing system.
        - Support agents can auto-approve refunds under $100.
        - For refunds exceeding $100, escalate to the Finance Manager for written approval.
        - Process the refund through Stripe or the original credit card gateway.
        - Standard processing takes 5-10 business days.
        """
        with open(DATA_DIR / "customer_refund_policy.txt", "w", encoding="utf-8") as f:
            f.write(refund_fallback.strip())

    # Document 3: Vendor XYZ project info
    vendor_content = """
    Department: Operations
    Date: 2026-05-10
    Title: Supplier Partnership: Vendor XYZ
    
    1. BACKGROUND
    Vendor XYZ is our primary logistics and supply chain provider. This document outlines active projects involving Vendor XYZ.
    
    2. PROJECTS INVOLVING VENDOR XYZ
    - Project Alpha: A supply chain automation program. Vendor XYZ handles container tracking integration. Managed by Sarah Connor.
    - Project Titan: Global warehouse expansion. Vendor XYZ provides logistics advice and heavy transport shipping.
    - Project Orion: Last-mile delivery optimization using drone delivery APIs. Vendor XYZ acts as the hardware consultant.
    
    3. COMPLIANCE & LEGAL
    A non-disclosure agreement (NDA) was signed between our company and Vendor XYZ on January 15, 2026. Under the agreement, all project specifications, api endpoints, and system architectures must remain confidential.
    
    4. CONTACTS
    For disputes or shipping delays, contact the Operations department or the Vendor XYZ accounts lead at accounts@vendorxyz.com.
    """
    with open(DATA_DIR / "vendor_xyz_partnership.txt", "w", encoding="utf-8") as f:
        f.write(vendor_content.strip())

if __name__ == "__main__":
    run_ingestion()
