import os
import re
from pathlib import Path
from datetime import datetime
import pypdf
import docx

def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text content from a PDF file using pypdf."""
    text_content = []
    try:
        reader = pypdf.PdfReader(str(file_path))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return "\n".join(text_content)

def extract_text_from_docx(file_path: Path) -> str:
    """Extract text content from a DOCX file using python-docx."""
    text_content = []
    try:
        doc = docx.Document(str(file_path))
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return "\n".join(text_content)

def extract_text_from_txt(file_path: Path) -> str:
    """Extract text content from a plain TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        return ""

def clean_extracted_text(text: str) -> str:
    """Normalize whitespace and merge single newlines within paragraphs."""
    if not text:
        return ""
    # Normalize carriage returns and horizontal whitespace
    text = text.replace('\r', '')
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Rebuild paragraphs to merge layout-induced line wraps
    lines = text.split('\n')
    paragraphs = []
    current_para = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_para:
                paragraphs.append(" ".join(current_para))
                current_para = []
        else:
            # Merge hyphenated words across lines
            if stripped.endswith('-') and len(stripped) > 1:
                current_para.append(stripped[:-1])
            else:
                current_para.append(stripped)
                
    if current_para:
        paragraphs.append(" ".join(current_para))
        
    cleaned = "\n\n".join(paragraphs)
    return re.sub(r' +', ' ', cleaned)

def extract_document_text(file_path: Path) -> str:
    """Detect file type and extract text content."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        text = extract_text_from_docx(file_path)
    elif suffix == ".txt":
        text = extract_text_from_txt(file_path)
    else:
        print(f"Unsupported file format: {suffix} for {file_path}")
        return ""
        
    return clean_extracted_text(text)

def detect_department(text: str, file_path: Path) -> str:
    """
    Heuristically detect department from text content or path names.
    Looks for department names in the file hierarchy first, then checks text.
    """
    # Check parent directory name first (e.g. data/HR/policy.txt -> HR)
    parent_name = file_path.parent.name
    if parent_name.lower() not in ["data", "rag", "."]:
        return parent_name.capitalize()

    # Scan first 2000 characters for department hints
    snippet = text[:2000].lower()
    
    # Check for direct declarations, e.g., "Department: Human Resources"
    match = re.search(r"(?:department|dept\.|team):\s*([a-zA-Z\s\-]+)", snippet)
    if match:
        dept = match.group(1).strip().title()
        if len(dept) < 30:  # sanity check length
            return dept
            
    # Keywords mapping
    keywords = {
        "Human Resources": ["hr", "human resources", "onboarding", "employee", "leave policy", "vacation", "recruiting", "benefits"],
        "Finance": ["finance", "refund", "billing", "invoice", "payment", "revenue", "audit", "budget", "tax"],
        "Legal": ["legal", "nda", "contract", "agreement", "compliance", "terms of service", "privacy policy", "litigation"],
        "Engineering": ["engineering", "software", "git", "code", "architecture", "deployment", "database", "api", "release notes"],
        "Operations": ["operations", "ops", "logistics", "vendor", "facility", "inventory", "procurement", "supply chain"],
        "Security": ["security", "firewall", "cybersecurity", "password", "gdpr", "data retention", "access control", "iso 27001"],
        "Marketing": ["marketing", "seo", "campaign", "social media", "brand", "advertising", "pr"]
    }
    
    for dept, terms in keywords.items():
        for term in terms:
            if re.search(r"\b" + re.escape(term) + r"\b", snippet):
                return dept
                
    return "General"

def extract_metadata(file_path: Path, text: str) -> dict:
    """Extract metadata such as filename, modified date, and department."""
    stats = file_path.stat()
    mod_time = datetime.fromtimestamp(stats.st_mtime)
    
    # Try to parse date from document text if possible (e.g., "Date: 2026-05-12")
    doc_date = None
    date_match = re.search(r"date:\s*([a-zA-Z0-9\s,\-]+)", text[:1000], re.IGNORECASE)
    if date_match:
        try:
            # Try to parse with standard formats
            date_str = date_match.group(1).strip()
            # Simple sanitization
            date_str_clean = re.sub(r'\s+', ' ', date_str)
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%B %d, %Y", "%d %B %Y"):
                try:
                    doc_date = datetime.strptime(date_str_clean, fmt)
                    break
                except ValueError:
                    continue
        except Exception:
            pass
            
    final_date = doc_date if doc_date else mod_time

    return {
        "filename": file_path.name,
        "filepath": str(file_path.absolute().as_posix()),
        "last_modified": mod_time.strftime("%Y-%m-%d %H:%M:%S"),
        "doc_date": final_date.strftime("%Y-%m-%d"),
        "department": detect_department(text, file_path),
        "file_size_bytes": stats.st_size
    }
